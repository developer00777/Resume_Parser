import base64
import io
import logging

from pypdf import PdfReader
from docx import Document
from fastapi import UploadFile, HTTPException

from app.config import settings

logger = logging.getLogger(__name__)

ALLOWED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}

# Minimum characters from PyPDF extraction to consider the PDF text-based.
# Below this threshold the PDF is treated as image-based and OCR is triggered.
_OCR_THRESHOLD = 100


def validate_file(file: UploadFile) -> str:
    """Validate uploaded file type. Returns file type string ('pdf' or 'docx')."""
    content_type = file.content_type or ""
    if content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {content_type}. Only PDF and DOCX are accepted.",
        )
    return ALLOWED_CONTENT_TYPES[content_type]


async def extract_text(file: UploadFile) -> str:
    """Extract text from uploaded PDF or DOCX file."""
    file_type = validate_file(file)
    content = await file.read()

    if len(content) > settings.max_file_size:
        raise HTTPException(
            status_code=413,
            detail=f"File size exceeds maximum allowed size of {settings.max_file_size} bytes.",
        )

    if file_type == "pdf":
        return await _extract_pdf(content)
    return _extract_docx(content)


async def validate_file_bytes(filename: str, content: bytes, content_type: str) -> str:
    """Extract text from raw bytes (used by background job tasks)."""
    if content_type not in ALLOWED_CONTENT_TYPES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {content_type}. Only PDF and DOCX are accepted.",
        )
    if len(content) > settings.max_file_size:
        raise HTTPException(
            status_code=413,
            detail=f"File '{filename}' exceeds maximum allowed size of {settings.max_file_size} bytes.",
        )
    file_type = ALLOWED_CONTENT_TYPES[content_type]
    if file_type == "pdf":
        return await _extract_pdf(content)
    return _extract_docx(content)


async def _ocr_pdf_via_vision(content: bytes) -> str:
    """
    OCR fallback for image-based PDFs.

    Sends the raw PDF as a base64 data URL to gpt-4o-mini via OpenRouter.
    No system dependencies required — reuses the existing OpenRouter HTTP client.
    Triggered only when PyPDF extracts fewer than _OCR_THRESHOLD characters.
    """
    from app.services.llm import _get_client

    logger.info("OCR fallback triggered — PDF appears image-based, sending to vision model")

    b64 = base64.b64encode(content).decode()
    data_url = f"data:application/pdf;base64,{b64}"

    client = _get_client()
    try:
        response = await client.post(
            "/chat/completions",
            json={
                "model": settings.openrouter_model,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": (
                                    "This is a scanned resume PDF. Extract ALL text exactly as it appears. "
                                    "Preserve names, dates, companies, job titles, skills, education, and all other details. "
                                    "Output plain text only — no commentary, no markdown, no JSON."
                                ),
                            },
                            {
                                "type": "file",
                                "file": {"url": data_url},
                            },
                        ],
                    }
                ],
                "max_tokens": 3000,
                "temperature": 0.0,
            },
        )
        response.raise_for_status()
    except Exception as e:
        logger.error(f"OCR vision call failed: {e}")
        raise HTTPException(
            status_code=422,
            detail="Could not extract text from PDF. The file appears image-based and OCR also failed.",
        )

    text = response.json()["choices"][0]["message"]["content"]
    logger.info(f"OCR extracted {len(text)} chars from image-based PDF")
    return text


async def _extract_pdf(content: bytes) -> str:
    """
    Extract text from PDF bytes.

    Strategy:
    1. PyPDF text extraction (fast, free, works for all text-based PDFs).
    2. If extracted text is below _OCR_THRESHOLD chars, the PDF is likely
       image-based — fall back to gpt-4o-mini vision OCR via OpenRouter.
    """
    try:
        reader = PdfReader(io.BytesIO(content))
        pages = []
        for page in reader.pages:
            # Use simple extraction — layout mode creates character-level spacing
            # artifacts on many PDF fonts (e.g. "Sa lesforce" instead of "Salesforce").
            text = page.extract_text()
            if text and text.strip():
                pages.append(text.strip())

        text = "\n\n".join(pages)

        if len(text.strip()) < _OCR_THRESHOLD:
            # Image-based PDF — PyPDF found no usable text layer.
            logger.warning(
                f"PyPDF extracted only {len(text.strip())} chars — triggering OCR fallback"
            )
            return await _ocr_pdf_via_vision(content)

        return text
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise HTTPException(status_code=422, detail=f"Failed to parse PDF: {e}")


def _extract_docx(content: bytes) -> str:
    """
    Extract text from DOCX bytes.

    Strategy:
    - Extracts body paragraphs (main document flow).
    - Also extracts text from tables (cell by cell) — common in resume templates.
    - Extracts text from text boxes / shapes via XML (for sidebar-style layouts).
    """
    try:
        doc = Document(io.BytesIO(content))
        parts: list[str] = []

        # Body paragraphs
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                parts.append(stripped)

        # Tables — many resume templates use tables for layout
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    parts.append(" | ".join(row_cells))

        # Text boxes / shapes (stored in body XML as <w:txbxContent>)
        try:
            import re as _re
            from docx.oxml.ns import qn
            body = doc.element.body
            for txbx in body.iter(qn('w:txbxContent')):
                for p in txbx.iter(qn('w:p')):
                    texts = [r.text for r in p.iter(qn('w:t')) if r.text]
                    combined = "".join(texts).strip()
                    if combined and combined not in parts:
                        parts.append(combined)
        except Exception:
            pass  # Text-box extraction is best-effort

        text = "\n".join(parts).strip()
        if not text:
            raise HTTPException(
                status_code=422,
                detail="Could not extract text from DOCX. The file may be empty.",
            )
        return text
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise HTTPException(status_code=422, detail=f"Failed to parse DOCX: {e}")
